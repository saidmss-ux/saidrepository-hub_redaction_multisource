"""
Document parsers for multi-format content extraction.

Each parser is responsible for extracting text from a specific file format.
Supports PDF, DOCX, and TXT with error handling and encoding detection.
"""

import logging
import os
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Dict, Optional

logger = logging.getLogger(__name__)


class ParserError(Exception):
    """Raised when parsing fails."""

    pass


class Parser(ABC):
    """Abstract base class for document parsers."""

    @abstractmethod
    def parse(self, file_path: str) -> Dict[str, str]:
        """
        Parse a file and extract text content.

        Args:
            file_path: Full path to the file to parse.

        Returns:
            Dictionary with keys:
            - text: Extracted text content
            - metadata: Optional metadata (page count, language, etc.)

        Raises:
            ParserError: If parsing fails.
        """
        pass

    @abstractmethod
    def supports(self, file_type: str) -> bool:
        """Check if this parser supports the given file type."""
        pass


class TXTParser(Parser):
    """Parser for plain text files with encoding detection."""

    SUPPORTED_TYPES = {"txt", "text"}

    def supports(self, file_type: str) -> bool:
        return file_type.lower() in self.SUPPORTED_TYPES

    def parse(self, file_path: str) -> Dict[str, str]:
        """
        Parse a text file with automatic encoding detection.

        Tries UTF-8 first, then falls back to other common encodings.
        """
        if not os.path.exists(file_path):
            raise ParserError(f"File not found: {file_path}")

        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252", "ascii"]
        last_error = None

        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    text = f.read()
                logger.info(f"Successfully parsed TXT with encoding: {encoding}")
                return {
                    "text": text,
                    "metadata": {"encoding": encoding, "file_type": "txt"},
                }
            except (UnicodeDecodeError, LookupError) as exc:
                last_error = exc
                continue

        raise ParserError(
            f"Could not parse TXT file with any supported encoding. Last error: {last_error}"
        )


class PDFParser(Parser):
    """Parser for PDF files using pdfplumber."""

    SUPPORTED_TYPES = {"pdf"}

    def supports(self, file_type: str) -> bool:
        return file_type.lower() in self.SUPPORTED_TYPES

    def parse(self, file_path: str) -> Dict[str, str]:
        """
        Parse a PDF file and extract text from all pages.

        Note: pdfplumber is only imported if actually used to avoid hard dependency.
        """
        if not os.path.exists(file_path):
            raise ParserError(f"File not found: {file_path}")

        try:
            import pdfplumber
        except ImportError:
            raise ParserError(
                "pdfplumber not installed. Install with: pip install pdfplumber"
            )

        try:
            text_parts = []
            metadata = {"page_count": 0, "file_type": "pdf"}

            with pdfplumber.open(file_path) as pdf:
                metadata["page_count"] = len(pdf.pages)

                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        text = page.extract_text() or ""
                        if text:
                            text_parts.append(f"[Page {page_num}]\n{text}")
                    except Exception as exc:
                        logger.warning(f"Failed to extract text from page {page_num}: {exc}")

            combined_text = "\n\n".join(text_parts) if text_parts else ""

            if not combined_text:
                logger.warning(f"PDF {file_path} contains no extractable text")

            return {"text": combined_text, "metadata": metadata}

        except Exception as exc:
            raise ParserError(f"Failed to parse PDF: {exc}")


class DOCXParser(Parser):
    """Parser for DOCX (Microsoft Word) files using python-docx."""

    SUPPORTED_TYPES = {"docx", "doc"}

    def supports(self, file_type: str) -> bool:
        return file_type.lower() in self.SUPPORTED_TYPES

    def parse(self, file_path: str) -> Dict[str, str]:
        """
        Parse a DOCX file and extract text from paragraphs and tables.

        Note: python-docx is only imported if actually used to avoid hard dependency.
        """
        if not os.path.exists(file_path):
            raise ParserError(f"File not found: {file_path}")

        try:
            from docx import Document
        except ImportError:
            raise ParserError(
                "python-docx not installed. Install with: pip install python-docx"
            )

        try:
            doc = Document(file_path)
            text_parts = []
            metadata = {"paragraphs": 0, "tables": 0, "file_type": "docx"}

            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            metadata["paragraphs"] = len(doc.paragraphs)

            # Extract text from tables
            for table_idx, table in enumerate(doc.tables, 1):
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    text_parts.append(f"\n[Table {table_idx}]\n" + "\n".join(table_text))
            metadata["tables"] = len(doc.tables)

            combined_text = "\n\n".join(text_parts) if text_parts else ""

            if not combined_text:
                logger.warning(f"DOCX {file_path} contains no extractable text")

            return {"text": combined_text, "metadata": metadata}

        except Exception as exc:
            raise ParserError(f"Failed to parse DOCX: {exc}")


class VideoParser(Parser):
    """Parser for video files using OpenAI Whisper."""

    SUPPORTED_TYPES = {"mp4", "webm", "mp3", "wav", "m4a"}

    def supports(self, file_type: str) -> bool:
        return file_type.lower() in self.SUPPORTED_TYPES

    def parse(self, file_path: str) -> Dict[str, str]:
        """
        Parse a video/audio file and transcribe to text using Whisper.

        Note: Requires 'openai-whisper' or 'faster-whisper' package.
        First call will download the model (~1-3 GB).
        """
        if not os.path.exists(file_path):
            raise ParserError(f"File not found: {file_path}")

        try:
            # Try faster-whisper first (lighter weight), fall back to openai-whisper
            try:
                from faster_whisper import WhisperModel
                model = WhisperModel("base")
                segments, info = model.transcribe(file_path)
                transcript = "\n".join([segment.text for segment in segments])
            except ImportError:
                try:
                    import whisper
                    model = whisper.load_model("base")
                    result = model.transcribe(file_path)
                    transcript = result.get("text", "")
                except ImportError:
                    raise ParserError(
                        "Whisper not installed. Install with: "
                        "pip install openai-whisper (or pip install faster-whisper)"
                    )
        except Exception as exc:
            raise ParserError(f"Failed to transcribe video: {exc}")

        metadata = {
            "file_type": "video",
            "transcribed": True,
            "model": "whisper-base",
        }

        return {"text": transcript, "metadata": metadata}


class ParserFactory:
    """Factory for selecting the appropriate parser based on file type."""

    def __init__(self):
        self.parsers = [PDFParser(), DOCXParser(), TXTParser(), VideoParser()]

    def get_parser(self, file_type: str) -> Optional[Parser]:
        """Get the appropriate parser for the file type."""
        file_type = file_type.lower().lstrip(".")
        for parser in self.parsers:
            if parser.supports(file_type):
                return parser
        return None

    def parse(self, file_path: str, file_type: str) -> Dict[str, str]:
        """
        Parse a file with automatic parser selection.

        Args:
            file_path: Path to the file to parse.
            file_type: File type (pdf, docx, txt, etc.).

        Returns:
            Parsed content as a dictionary.

        Raises:
            ParserError: If no parser supports the file type or parsing fails.
        """
        parser = self.get_parser(file_type)
        if not parser:
            raise ParserError(
                f"Unsupported file type: {file_type}. "
                f"Supported: {', '.join(['pdf', 'docx', 'txt'])}"
            )
        return parser.parse(file_path)
